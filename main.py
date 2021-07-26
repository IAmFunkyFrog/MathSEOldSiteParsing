# Заметка: как мне удалось выяснить, группы x7x - програмные инженеры, а группы x4x - матобесы

import os
import sys
from os.path import splitext
import requests
from bs4 import BeautifulSoup
from transliterate import translit
from pdfminer import high_level
import re
import json
import docx

SECRET_KEY = '8795e9a6d5199969b1f8e6385a15055b'
UPLOAD_URL = "https://se.math.spbu.ru/post_theses"
TEXT_PATH = "/media/stepan-trefilov/Share/report/text/"
SLIDES_PATH = "/media/stepan-trefilov/Share/report/slides/"
SUPERVISOR_REVIEW_PATH = "/media/stepan-trefilov/Share/report/review/"

SUPERVISORS = (
    'Кириленко',
    'Баклановский',
    'Литвинов',
    'Подкопаев',
    'Пименов',
    'Немешев',
    'Григорьев',
    'Булычев',
    'Монькин',
    'Лазарева',
    'Губанов',
    'Сартасов',
    'Брыксин',
    'Амелин',
    'Иноземцев',
    'Давыденко',
    'Невоструев',
    'Нестеров',
    'Вахитов',
    'Новиков',
    'Романовский',
    'Чурилин',
    'Коновалов',
    'Малов',
    'Маров',
    'Дымникова',
    'Чурилин',
    'Николенко',
    'Смирнов',
    'Терехов',
    'Лазарева',
    'Вояковская',
    'Праздников',
    'Граничин',
    'Мордвинов',
    'Боташ',
    'Шалымов',
    'Васильев',
    'Анисимов',
    'Пахомов',
    'Граничин',
    'Королёв',
    'Козлов',
    'Медведев',
    'Куралёнок',
    'Щитинин',
    'Платонов',
    'Оносовский',
    'Вяххи',
    'Семихатский',
    'Соломатов',
    'Данильченко',
    'Петров',
    'Зеленчук',
    'Николаев',
    'Абусалимов',
    'Белогрудов',
    'Суворов',
    'Минкин',
    'Бугайченко',
    'Полозов',
    'Осечкина',
    'Дыдычкин',
    'Козловский'
)

# Флаг скачки файлов с сайта
download = True

# Флаг загрузки файлов на сайт
UPLOAD_FLAG = True


def download_file(uri, safe_filename, save_path="./report/"):
    # Skip if download == false
    if not download:
        print("Download flag if False")
        return

    print(save_path + safe_filename)
    r = requests.get(uri, allow_redirects=True)
    open(save_path + safe_filename, 'wb').write(r.content)
    print("Downloaded " + safe_filename)


def rename_file(old_filename, new_filename):
    print("Renaming " + old_filename + " to " + new_filename)
    os.rename(old_filename, new_filename)


def get_supervisor_from_text(text):
    try:
        supervisor_re = re.search(r".{0,250}[Нн]аучный\sруководитель.{0,250}", text)[0]
    except TypeError:
        supervisor_re = text + ' '
        print("Error with parsing text")
    supervisor = ''
    print("String that must contain supervisor: " + supervisor_re)
    for supervisor_string in SUPERVISORS:
        if supervisor_re.find(supervisor_string + ' ', 0) > -1:
            supervisor = supervisor_string
            break
    return supervisor


def get_supervisor_from_file(path):
    text_extension = splitext(path)[-1]
    text_of_work = ''
    if text_extension.find('.pdf', 0) > -1:
        text_of_work = high_level.extract_text(path)
    elif text_extension.find('.doc', 0) > -1 or text_extension.find('.docx', 0) > -1:
        try:
            document = docx.Document(path)
        except KeyError:
            print("Error while parsing .doc")
            return ''
        except ValueError:
            print("Error while parsing .doc - it is not .doc or .docx")
            return ''
        except:
            print("Error while parsing .doc - Mb file is corrupted")
            return ''
        for paragraph in document.paragraphs:
            text_of_work = text_of_work + paragraph.text + ' '
    text_of_work = text_of_work.replace('\n', ' ')
    return get_supervisor_from_text(text_of_work)


def upload_on_site(thesis_info, text_filename, slides_filename='', supervisor_review_filename=''):
    print(thesis_info)

    if not UPLOAD_FLAG:
        print("Upload on site disabled")
        return

    report_text = TEXT_PATH + text_filename
    files = [
        ('thesis_text', (report_text, open(report_text, 'rb'), 'application/octet')),
        ('thesis_info', ('thesis_info', json.dumps(thesis_info), 'application/json')),
    ]

    if slides_filename != '':
        presentation = SLIDES_PATH + slides_filename
        files.append(('presentation', (presentation, open(presentation, 'rb'), 'application/octet')))

    if supervisor_review_filename != '':
        supervisor_review = SUPERVISOR_REVIEW_PATH + supervisor_review_filename
        files.append(('supervisor_review', (supervisor_review, open(supervisor_review, 'rb'), 'application/octet')))

    r = requests.post(UPLOAD_URL, files=files, allow_redirects=False)
    print(str(r.content, 'utf-8'))


def get_text_filename(author_en, year, extension):
    return author_en + "_Bachelor_Report_" + str(year) + "_text" + extension


def get_slides_filename(author_en, year, extension):
    return author_en + "_Bachelor_Report_" + str(year) + "_slides" + extension


def get_supervisor_review_filename(author_en, year, extension):
    return author_en + "_Bachelor_Report_" + str(year) + "_supervisor_review" + extension


def get_2017_reports():
    session = requests.session()
    url = 'https://oops.math.spbu.ru/SE/YearlyProjects/spring-2017'
    year = 2017

    response = session.get(url)

    if response.status_code != 200:
        print("Response status " + str(response.status_code))
        sys.exit(0)

    soup = BeautifulSoup(response.text, "lxml")

    # Найдем таблицы содержащие курсовые работы
    tables = soup.select(".listing")

    # Разберем первую таблицу, бакалавры 371 группы
    for row in tables[0].find_all('tr'):
        cols = row.find_all('td')
        if len(cols) != 3:
            print("Error while parsing cols in table")
            continue

        author = cols[0].text
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = cols[1].text

        # находим ссылки на текст слайды и рецензию
        anchors = cols[2].find_all('a')

        # Скачиваем текст
        text_uri = anchors[0].get('href')
        text_extension = splitext(text_uri)[1]
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Download text: " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Скачиваем слайды
        slides_uri = anchors[1].get('href')
        slides_extension = splitext(slides_uri)[1]
        slides_filename = get_slides_filename(author_en, year, slides_extension)
        print("Download slides: " + slides_filename)
        download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Скачиваем отзыв
        supervisor_review_uri = anchors[2].get('href')
        supervisor_review_extension = splitext(supervisor_review_uri)[1]
        supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
        print("Download supervisor review: " + supervisor_review_filename)
        download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Достаем имя научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue

        print("Supervisor: " + supervisor)

        # Генерируем метаинформацию и загружаем

        thesis_info = {'type_id': 2, 'course_id': 2, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

    # Разберем вторую таблицу, бакалавры 344 группы
    for row in tables[1].find_all('tr'):
        cols = row.find_all('td')
        if len(cols) != 4:
            print("Error while parsing cols in table")
            continue

        author = cols[0].text
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = cols[1].text

        # находим ссылки на текст слайды и рецензию
        anchors = cols[3].find_all('a')

        # Скачиваем текст
        text_uri = anchors[0].get('href')
        text_extension = splitext(text_uri)[1]
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Download text: " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Скачиваем слайды
        slides_uri = anchors[1].get('href')
        slides_extension = splitext(slides_uri)[1]
        slides_filename = get_slides_filename(author_en, year, slides_extension)
        print("Download slides: " + slides_filename)
        download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Скачиваем отзыв
        supervisor_review_filename = ''
        if len(anchors) > 2:
            supervisor_review_uri = anchors[2].get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Достаем имя научника
        supervisor = get_supervisor_from_text(cols[2].text)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue

        print("Supervisor: " + supervisor)

        # Генерируем метаинформацию и загружаем

        thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)


def get_2016_reports():
    session = requests.session()
    url = 'https://oops.math.spbu.ru/SE/YearlyProjects/spring-2016'
    year = 2016

    response = session.get(url)

    if response.status_code != 200:
        print("Response status " + str(response.status_code))
        sys.exit(0)

    soup = BeautifulSoup(response.text, "lxml")

    # Найдем таблицы содержащие курсовые работы
    tables = soup.select(".listing")

    # Разберем первую таблицу, магистры 546 группа
    for row in tables[0].find_all('tr'):
        cols = row.find_all('td')
        if len(cols) != 3:
            print("Error while parsing cols in table")
            continue

        author = cols[0].text
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = cols[1].text

        # находим ссылки на текст
        anchors = cols[2].find_all('a')

        # Скачиваем текст
        text_uri = anchors[0].get('href')
        text_extension = splitext(text_uri)[1]
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Download text: " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем имя научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue

        print("Supervisor: " + supervisor)

        report_text = "report/text/" + text_filename

        thesis_info = {'type_id': 2, 'course_id': 3, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename)

    # Разберем вторую таблицу, бакалавры 371 группа
    for row in tables[1].find_all('tr'):
        cols = row.find_all('td')
        if len(cols) != 3:
            print("Error while parsing cols in table")
            continue

        author = cols[0].text
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = cols[1].text

        # находим ссылки на текст слайды и рецензию
        anchors = cols[2].find_all('a')

        if len(anchors) == 0:
            print("error while finding anchor to report")
            continue

        # Скачиваем текст
        text_uri = anchors[0].get('href')
        text_extension = splitext(text_uri)[1]
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Download text: " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Скачиваем слайды
        slides_uri = anchors[1].get('href')
        slides_extension = splitext(slides_uri)[1]
        slides_filename = get_slides_filename(author_en, year, slides_extension)
        print("Download slides: " + slides_filename)
        download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Скачиваем отзыв
        supervisor_review_filename = ''
        if len(anchors) > 2:
            supervisor_review_uri = anchors[2].get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Достаем имя научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue

        print("Supervisor: " + supervisor)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 2, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

    # Разберем третью таблицу, бакалавры 344 группа
    for row in tables[2].find_all('tr'):
        cols = row.find_all('td')
        if len(cols) != 3:
            print("Error while parsing cols in table")
            continue

        author = cols[0].text
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = cols[1].text

        # находим ссылки на текст слайды и рецензию
        anchors = cols[2].find_all('a')

        if len(anchors) == 0:
            print("error while finding anchor to report")
            continue

        # Скачиваем текст
        text_uri = anchors[0].get('href')
        text_extension = splitext(text_uri)[1]
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Download text: " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Скачиваем слайды
        slides_uri = anchors[1].get('href')
        slides_extension = splitext(slides_uri)[1]
        slides_filename = get_slides_filename(author_en, year, slides_extension)
        print("Download slides: " + slides_filename)
        download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Скачиваем отзыв
        supervisor_review_filename = ''
        if len(anchors) > 2:
            supervisor_review_uri = anchors[2].get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Достаем имя научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue

        print("Supervisor: " + supervisor)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)


def get_2015_fall():
    session = requests.session()
    url = 'https://oops.math.spbu.ru/SE/YearlyProjects/autumn-2015/magistracy-564'
    year = 2015

    response = session.get(url)

    if response.status_code != 200:
        print("Response status " + str(response.status_code))
        sys.exit(0)

    soup = BeautifulSoup(response.text, "lxml")
    spans = soup.select('.summary')

    # Переберем всех магистров 546 группы
    for span in spans:
        anchor = span.find('a')
        splited_author_and_theme = anchor.text.split('. ')
        author = splited_author_and_theme[0]
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = splited_author_and_theme[1]

        # Находим текст
        response = session.get(anchor.get("href"))

        if response.status_code != 200:
            print("Response status " + str(response.status_code))
            sys.exit(0)

        text_soup = BeautifulSoup(response.text, "lxml")
        anchors_to_text = text_soup.find_all('a')
        text_filename = ''
        for a in anchors_to_text:
            if a.text.find(author, 0) > -1:
                text_extension = splitext(a.text)[1].replace('\n', '')
                text_filename = get_text_filename(author_en, year, text_extension)
                print("Download " + text_filename)
                download_file(a.get('href'), text_filename, TEXT_PATH)
                break

        if text_filename == '':
            print("Error while parsing text filename")
            continue

        # Находим имя научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue

        print("Supervisor: " + supervisor)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 3, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename)


def get_2015_spring():
    session = requests.session()
    url = 'https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/list'
    year = 2015

    response = session.get(url)

    if response.status_code != 200:
        print("Response status " + str(response.status_code))
        sys.exit(0)

    soup = BeautifulSoup(response.text, "lxml")
    uls = soup.find('h3', text='344 группа').parent.find_all('ul')

    # Парсинг 344 группы
    for li in uls[0].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')
        # Достаем научника
        supervisor = get_supervisor_from_text(li.text)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        text_of_work = ''
        if text_extension.find('.pdf', 0) > -1:
            text_of_work = high_level.extract_text(TEXT_PATH + text_tmp_name)
        elif text_extension.find('.doc', 0) > -1 or text_extension.find('.docx', 0) > -1:
            document = docx.Document(TEXT_PATH + text_tmp_name)
            for paragraph in document.paragraphs:
                text_of_work = text_of_work + paragraph.text + ' '
        text_of_work.replace('\n', ' ')
        author_re = re.search(re.compile(surname + "\s+\w+\s+\w+\s"), text_of_work)
        if author_re is None:
            print("Error while parsing author name")
            continue

        author = author_re[0].replace('\n', '')
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(surname + "\s+(.+)\("), li.text)[1]

        print("Removing tmp text " + text_tmp_name)
        os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

    # Парсинг 371 группы
    for li in uls[1].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')
        # Достаем научника
        supervisor = get_supervisor_from_text(li.text)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        text_of_work = ''
        if text_extension.find('.pdf', 0) > -1:
            text_of_work = high_level.extract_text(TEXT_PATH + text_tmp_name)
        elif text_extension.find('.doc', 0) > -1 or text_extension.find('.docx', 0) > -1:
            document = docx.Document(TEXT_PATH + text_tmp_name)
            for paragraph in document.paragraphs:
                text_of_work = text_of_work + paragraph.text + ' '
        text_of_work.replace('\n', ' ')
        author_re = re.search(re.compile(surname + "\s+\w+\s+\w+\s"), text_of_work)
        if author_re is None:
            print("Error while parsing author name")
            continue

        author = author_re[0].replace('\n', '')
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(surname + "\s+(.+)\("), li.text)[1]

        print("Removing tmp text " + text_tmp_name)
        os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 2, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)


def get_2014():
    session = requests.session()
    url = 'https://oops.math.spbu.ru/SE/YearlyProjects/2014/list'
    year = 2014

    response = session.get(url)

    if response.status_code != 200:
        print("Response status " + str(response.status_code))
        sys.exit(0)

    soup = BeautifulSoup(response.text, "lxml")
    uls = soup.find('h3', text='344 группа').parent.find_all('ul')

    # Парсинг 344 группы
    for li in uls[0].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')
        # Достаем научника
        supervisor = get_supervisor_from_text(re.search(r"\((.+)\)", li.text)[1])

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        author = li.find('b').text.lstrip()
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(author + "\s+(.+)\("), li.text)[1]

        print("Removing tmp text " + text_tmp_name)
        if download:
            os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

    # Парсинг 371 группы
    for li in uls[1].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')
        # Достаем научника
        supervisor = get_supervisor_from_text(re.search(r"\((.+)\)", li.text)[1])

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        author = li.find('b').text.lstrip()
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(author + "\s+(.+)\("), li.text)[1]

        print("Removing tmp text " + text_tmp_name)
        if download:
            os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 2, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

    # Парсинг 444 группы
    for li in uls[2].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')
        # Достаем научника
        supervisor = get_supervisor_from_text(re.search(r"\((.+)\)", li.text)[1])

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        author = li.find('b').text.lstrip()
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(author + "\s+(.+)\("), li.text)[1]

        print("Removing tmp text " + text_tmp_name)
        if download:
            os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)


def get_2013():
    session = requests.session()
    url = 'https://oops.math.spbu.ru/SE/YearlyProjects/2013/list'
    year = 2013

    response = session.get(url)

    if response.status_code != 200:
        print("Response status " + str(response.status_code))
        sys.exit(0)

    soup = BeautifulSoup(response.text, "lxml")
    uls = soup.find('h3', text='341 группа').parent.find_all('ul')

    # Парсинг 341 группы
    for li in uls[0].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')

        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        author = li.find('b').text.lstrip()
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(author + "\s+(.+)"), li.text)[1]
        print("Work name " + name)

        print("Removing tmp text " + text_tmp_name)
        if download:
            os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        if supervisor_review_anchor is None:
            supervisor_review_anchor = li.find('a', text='Отзыв (часть 1)')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

        # Парсинг 344 группы
        for li in uls[1].find_all('li'):
            surname = li.find('b').text.split('.')[-1].replace(' ', '')

            anchors = li.find_all('a')
            # Достаем текст
            text_anchor = li.find('a', text='Отчёт')
            text_uri = text_anchor.get('href')
            text_extension = splitext(text_uri)[-1]
            text_tmp_name = get_text_filename(surname, year, text_extension)
            print("Download tmp text: " + text_tmp_name)
            download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

            # Достаем имя студента из текста
            author = li.find('b').text.lstrip()
            author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
            name = re.search(re.compile(author + "\s+(.+)"), li.text)[1]
            print("Work name " + name)

            print("Removing tmp text " + text_tmp_name)
            if download:
                os.remove(TEXT_PATH + text_tmp_name)
            text_filename = get_text_filename(author_en, year, text_extension)
            print("Downloading text " + text_filename)
            download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

            # Достаем научника
            supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

            if supervisor == '':
                print("Error while parsing supervisor")
                continue
            print("Supervisor " + supervisor)

            # Достаем слайды
            slides_anchor = li.find('a', text='Презентация')
            slides_filename = ''
            if slides_anchor is not None:
                slides_uri = slides_anchor.get('href')
                slides_extension = splitext(slides_uri)[1]
                slides_filename = get_slides_filename(author_en, year, slides_extension)
                print("Download slides: " + slides_filename)
                download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

            # Достаем отзыв научника
            supervisor_review_anchor = li.find('a', text='Отзыв')
            supervisor_review_filename = ''
            if supervisor_review_anchor is not None:
                supervisor_review_uri = supervisor_review_anchor.get('href')
                supervisor_review_extension = splitext(supervisor_review_uri)[1]
                supervisor_review_filename = get_supervisor_review_filename(author_en, year,
                                                                            supervisor_review_extension)
                print("Download supervisor review: " + supervisor_review_filename)
                download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

            # Генерируем метаинформацию и загружаем
            thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                           'supervisor': supervisor, 'publish_year': year,
                           'secret_key': SECRET_KEY}

            upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

    # Парсинг 361 группы
    for li in uls[2].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')
        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        author = li.find('b').text.lstrip()
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(author + "\s+(.+)"), li.text)[1]
        print("Work name " + name)

        print("Removing tmp text " + text_tmp_name)
        if download:
            os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        # TODO: попросить у зеленчука добавить 361 группу в бд и поменять course_id
        thesis_info = {'type_id': 2, 'course_id': 5, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)

    # Парсинг 445 группы
    for li in uls[3].find_all('li'):
        surname = li.find('b').text.split('.')[-1].replace(' ', '')

        anchors = li.find_all('a')
        # Достаем текст
        text_anchor = li.find('a', text='Отчёт')
        text_uri = text_anchor.get('href')
        text_extension = splitext(text_uri)[-1]
        text_tmp_name = get_text_filename(surname, year, text_extension)
        print("Download tmp text: " + text_tmp_name)
        download_file(url + "/" + text_uri, text_tmp_name, TEXT_PATH)

        # Достаем имя студента из текста
        author = li.find('b').text.lstrip()
        author_en = translit(author, 'ru', reversed=True).replace(" ", "_")
        name = re.search(re.compile(author + "\s+(.+)"), li.text)[1]
        print("Work name " + name)

        print("Removing tmp text " + text_tmp_name)
        if download:
            os.remove(TEXT_PATH + text_tmp_name)
        text_filename = get_text_filename(author_en, year, text_extension)
        print("Downloading text " + text_filename)
        download_file(url + "/" + text_uri, text_filename, TEXT_PATH)

        # Достаем научника
        supervisor = get_supervisor_from_file(TEXT_PATH + text_filename)

        if supervisor == '':
            print("Error while parsing supervisor")
            continue
        print("Supervisor " + supervisor)

        # Достаем слайды
        slides_anchor = li.find('a', text='Презентация')
        slides_filename = ''
        if slides_anchor is not None:
            slides_uri = slides_anchor.get('href')
            slides_extension = splitext(slides_uri)[1]
            slides_filename = get_slides_filename(author_en, year, slides_extension)
            print("Download slides: " + slides_filename)
            download_file(url + "/" + slides_uri, slides_filename, SLIDES_PATH)

        # Достаем отзыв научника
        supervisor_review_anchor = li.find('a', text='Отзыв')
        supervisor_review_filename = ''
        if supervisor_review_anchor is not None:
            supervisor_review_uri = supervisor_review_anchor.get('href')
            supervisor_review_extension = splitext(supervisor_review_uri)[1]
            supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
            print("Download supervisor review: " + supervisor_review_filename)
            download_file(url + "/" + supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

        # Генерируем метаинформацию и загружаем
        thesis_info = {'type_id': 2, 'course_id': 1, 'name_ru': name, 'author': author,
                       'supervisor': supervisor, 'publish_year': year,
                       'secret_key': SECRET_KEY}

        upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)


def upload_one_report(thesis_info, text_uri, slides_uri='', supervisor_review_uri=''):
    if text_uri == '':
        print("Invalid text_uri")
        return

    year = thesis_info['publish_year']
    author = thesis_info['author']
    author_en = translit(author, 'ru', reversed=True).replace(" ", "_")

    text_extension = '.' + text_uri.split('.')[-1]
    if text_extension != '.pdf' and text_extension != '.doc' and text_extension != '.docx':
        text_extension = '.pdf'
    text_filename = get_text_filename(author_en, year, text_extension)
    download_file(text_uri, text_filename, TEXT_PATH)

    slides_filename = ''
    if slides_uri != '':
        slides_extension = '.' + slides_uri.split('.')[-1]
        print(slides_extension)
        slides_filename = get_slides_filename(author_en, year, slides_extension)
        download_file(slides_uri, slides_filename, SLIDES_PATH)
    supervisor_review_filename = ''
    if supervisor_review_uri != '':
        supervisor_review_extension = '.' + supervisor_review_uri.split('.')[-1]
        supervisor_review_filename = get_supervisor_review_filename(author_en, year, supervisor_review_extension)
        download_file(supervisor_review_uri, supervisor_review_filename, SUPERVISOR_REVIEW_PATH)

    upload_on_site(thesis_info, text_filename, slides_filename, supervisor_review_filename)


# скрипт загрузки проблемных работ, которые не удалось просто распарсить
def bruteforce_2012():
    # 2012 год
    # 341 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка 3D пиксельного движка',
         'author': 'Осипов Никита Алексеевич',
         'supervisor': 'Оносовский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/341_Osipov_report.doc',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/341_Osipov_review.pdf'
    )
    # 345 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Фреймворки юнит-тестирования для С++',
         'author': 'Бажутин Михаил Сергеевич',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Bazhutin_report.doc',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Bazhutin_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка 3D пиксельного графического движка',
         'author': 'Байцерова Юлия Сергеевна',
         'supervisor': 'Оносовский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Baytserova_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Baytserova_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Baytserova_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Терминальный Android-клиент для распределенных приложений на базе платформы Ubiq Mobile',
         'author': 'Бумаков Никита Вячеславович',
         'supervisor': 'Оносовский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Bumakov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Bumakov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Bumakov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Статический анализ кода языка Ruby',
         'author': 'Денисов Юрий Борисович',
         'supervisor': 'Ушаков', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Denison_report.doc',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Denison_review.png'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Визуальный язык задания ограничений на модели в QReal',
         'author': 'Дерипаска Анна Олеговна',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Deripaska_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Deripaska_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Deripaska_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Структура хранения индексов закэшированных страйпов и ссылок на них',
         'author': 'Дудин виктор Дмитриевич',
         'supervisor': 'Короткевич', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Dudin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Dudin_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Dudin_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка средства проверки корректности адресов возврата на платформе S2E',
         'author': 'Евард Вадим Евгеньевич',
         'supervisor': 'Зеленчук', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Evard_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Evard_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Evard_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация механизмов виртуальной памяти для x86 архитектуры в ОСРВ Embox',
         'author': 'Ефимов Глеб Дмитриевич',
         'supervisor': 'Бондарев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Efimov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Efimov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Efimov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Определение расстояния между точкой и множеством, представленным бинарной диаграммой решений',
         'author': 'Зубаревич Дмитрий Александрович',
         'supervisor': 'Бугайченко', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Zubarevich_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Zubarevich_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Zubarevich_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Алгоритмы расчета RAID 6',
         'author': 'Калмук Александр Игоревич',
         'supervisor': 'Короткевич', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kalmuk_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kalmuk_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Инструмент для оценки алгоритмов дедупликации',
         'author': 'Кладов Алексей Александрович',
         'supervisor': 'Луцив', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kladov_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kladov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Трассировки ОСРВ Embox',
         'author': 'Крамар Алина Сергеевна',
         'supervisor': 'Бондарев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kramar_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kramar_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Поддержка механизма рефакторингов в metaCASE-системе QReal',
         'author': 'Кузенкова Анастасия Сергеевна',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kuzenkova_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Kuzenkova_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация алгоритмов расчета RAID 6 с использованием встроенных функций SSE',
         'author': 'Макулов Рустам Наилевич',
         'supervisor': 'Короткевич', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Makulov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Makulov_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Makulov_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Переиспользование кода в визуальных языках программирования',
         'author': 'Нефёдов Ефим Андреевич',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Nefedov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Nefedov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Nefedov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Исследование и тестирование семплирующего метода профайлинга на примере профилировщика производительности Intel VTune Amplifier XE 2011',
         'author': 'Одеров Роман Сергеевич',
         'supervisor': 'Баклановский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Oderov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Oderov_presentation.pptx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Средства описания генераторов кода для предметно-ориентированных решений в metaCASE-средстве QReal',
         'author': 'Подкопаев Антон Викторович',
         'supervisor': 'Брыксин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Podkopaev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Podkopaev_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Podkopaev_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Интерпретация метамоделей в metaCASE-системе QReal',
         'author': 'Птахина Алина Ивановна',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Ptakhina_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Ptakhina_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Ptakhina_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Модуль сбора информации о производительности процессоров Intel с использованием PMU для профайлера ядра ОС MS WS2008R2',
         'author': 'Серко Сергей Анатольевич',
         'supervisor': 'Баклановский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Serko_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Serko_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Юзабилити в проекте QReal:Robots',
         'author': 'Соковикова Наталья Алексеевна',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Sokovikova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Sokovikova_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Sokovikova_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Исследование эффективности дедупликации с использованием цепочек преобразований при помощи разработанного инструментального средства',
         'author': 'Соса Укатерина Андреевна',
         'supervisor': 'Луцив', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Sosa_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Sosa_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Архитектура и прототипирование metaCASE-системы',
         'author': 'Таран Кирилл Сергеевич',
         'supervisor': 'Брыксин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Taran_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Taran_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Taran_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Алгоритмы расчёта RAID 6',
         'author': 'Тюшев Кирилл Игоревич',
         'supervisor': 'Короткевич', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Tyushev_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Tyushev_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация уровня изоляции Read Committed для древовидных структур данных',
         'author': 'Федотовский Павел Валерьевич',
         'supervisor': 'Чернышев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Fedotovskij_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Fedotovskij_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Fedotovskij_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Выделение научных сообществ на основе анализа библиографических данных',
         'author': 'Филатов Владимир Константинович',
         'supervisor': 'Суворов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Filatov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Ibragimov_Filatov_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Сравнение распределителей памяти для многопоточного обработчика транзакций',
         'author': 'Чередник Кирилл Евгеньевич',
         'supervisor': 'Смирнов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Cherednik_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/345/345_Cherednik_review.doc'
    )
    # 361 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Поисковые алгоритмы на блоке графического процессора',
         'author': 'Алексеев Илья Владимирович',
         'supervisor': 'Губанов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Alekseev_report.docx',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Разработка редактора диаграмм для облачной технологии создания мобильных приложений',
         'author': 'Белокуров Дмитрий Николаевич',
         'supervisor': 'Брыксин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Belokurov_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Belokurov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Обзор реализации механизма циклической разработки диаграмм классов и программного кода в современных UML-средствах',
         'author': 'Бусыгина Мария',
         'supervisor': 'Кознов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Busygina_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Стабилизация показаний G-сенсора при помощи камеры',
         'author': 'Говейнович Сергей Геннадьевич',
         'supervisor': 'Оносовский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Goveynovich_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Goveynovich_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Разработка сети простых вычислительных процессоров и фильтра в рамках студенческого проекта МПВ',
         'author': ' Забранский Дмитрий',
         'supervisor': 'Кривошеин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Zabranskiy_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Zabranskiy_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Разработка трехмерного игрового движка для игры под Android',
         'author': 'Зольников Павел Евгеньевич',
         'supervisor': 'Оносовский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Zolnikov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Zolnikov_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Zolnikov_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Создание подсистемы управления рисками: разработка бизнес-логики подсистемы',
         'author': 'Зубрилин',
         'supervisor': 'Кияев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Zubrilin_Report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Выделение научных сообществ на основе анализа библиографических данных',
         'author': 'Ибрагимов Рустам',
         'supervisor': 'Суворов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Ibragimov_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Анализ тональности текста',
         'author': 'Калмыков Алексей Владимирович',
         'supervisor': 'Губанов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Kalmykov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Kalmykov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Kalmykov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Создание информационной подсистемы для распределенной подготовки стартапа',
         'author': 'Калугин',
         'supervisor': 'Кияев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Kalugin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Kalugin_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': ' Реализация эвристик языка ДРАКОН в metaCASE-средстве QReal',
         'author': 'Колантаевская Анна Сергеевна',
         'supervisor': 'Брыксин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Kolantaevskaya_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Kolantaevskaya_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': ' Тиражирование мобильных приложений для платформы iOS и Android',
         'author': 'Коршаков Степан Андреевич',
         'supervisor': 'Сабашный', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Korshakov_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Korshakov_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Оценка сайта на наличие нежелательного контента',
         'author': 'Кривых Алексей',
         'supervisor': 'Тарасов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Krivykh_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Krivykh_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Разработка интерфейса программирования приложений и добавление скриптовой функциональности для ПО криминалистического анализа',
         'author': 'Макеев Давид Александрович',
         'supervisor': 'Губанов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Makeev_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Makeev_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Выделение человека на изображении',
         'author': 'Монькин Александр Александрович',
         'supervisor': 'Петров', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Monkin_S_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Monkin_S_review.jpeg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Вычисление объема трехмерного объекта в задаче планирования хирургической операции',
         'author': 'Монькин Сергей Александрович',
         'supervisor': 'Петров', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Monkin_S_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Monkin_S_review.jpeg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Трехмерная модель робота в QReal:Robots',
         'author': 'Павлов Сергей Николаевич',
         'supervisor': 'Брыксин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Pavlov_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Выделение групп пользователей в социальных сетях',
         'author': 'Никита Симонов',
         'supervisor': 'Суворов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Simonov_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Реализация модулей ввода/вывода ПВП в связке с ядром DSP48E в рамках проекта МПВ',
         'author': 'Солдатов Дмитрий',
         'supervisor': 'Кривошеин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Soldatov_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Soldatov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Верификация дизассемблера x86-64',
         'author': 'Тенсин Егор Дмитриевич',
         'supervisor': 'Баклановский', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Tensin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Tensin_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Разработка модуля памяти для многоядерного потокового вычислителя',
         'author': 'Тодорук',
         'supervisor': 'Кривошеин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Todoruk_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Todoruk_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': ' Транслятор микрокода для многоядерного потокового вычислителя',
         'author': 'Улитин Александр',
         'supervisor': 'Кривошеин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Ulitin_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Ulitin_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Создание подсистемы управления рисками: разработка архитектуры подсистемы',
         'author': 'Яськов',
         'supervisor': 'Кияев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/361/361_Yaskov_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Создание генератора GLR трансляторов для .NET',
         'author': 'Авдюхин Дмитрий Алексеевич',
         'supervisor': 'Кириленко', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Avdyukhin_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Avdyukhin_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Организация надежных соединений через виртуальные каналы',
         'author': 'Бурдун Егор Федорович',
         'supervisor': 'Бондарев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Burdun_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Burdun_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Алгоритм построения оценок весов интентов для многозначных запросов',
         'author': 'Григорьев Артем Валерьевич',
         'supervisor': 'Грауэр', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Grigoriev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Grigoriev_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Grigoriev_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Обучение информатике в школах и ВУЗах на примере ОСРВ Embox',
         'author': 'Дзендик Дарья Анатольевна',
         'supervisor': 'Бондарев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Dzendzik_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Dzendzik_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Dzendzik_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Распознавание языка жестов на видео потоке',
         'author': 'Землянская Светлана Андреевна',
         'supervisor': 'Граничин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Zemlyanskaya_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Zemlyanskaya_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Zemlyanskaya_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Генерация кода для платформы Ubiq Mobile',
         'author': 'Иванов Всеволод Юрьевич',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Ivanov_report.docx',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Ivanov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Программная платформа для встраиваемых решений',
         'author': 'Козлов Антон Павлович',
         'supervisor': 'Бондарев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Kozlov_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Kozlov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Повышение прозрачности сайта госзакупок РФ',
         'author': 'Коноплев Юрий',
         'supervisor': 'Сысоев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Konoplev_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Konoplev_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Оптимизация вычислений за счет эффективных структур данных в ОС Embox',
         'author': 'Мальковский Николай Владимирович',
         'supervisor': 'Бондарев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Malkovsky_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Malkovsky_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Malkovsky_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Механизм автоматической генерации мигратора базы данных информационной системы при изменениях модели предметной области',
         'author': 'Михайлов Дмитрий Петрович',
         'supervisor': 'Нестеров', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Mikhaylov_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Mikhaylov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация настраиваемого графического представления элемента на диаграмме в QReal',
         'author': 'Мордвинов Дмитрий Александрович',
         'supervisor': 'Брыксин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Mordvinov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Mordvinov_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Mordvinov_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Реализация алгоритма Semi-Global Matching',
         'author': 'Мокаев Руслан',
         'supervisor': 'Пименов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Mokaev_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Статическая верификация для языка HaSCoL',
         'author': 'Найданов Дмитрий Геннадьевич',
         'supervisor': 'Медведев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Naydanov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Naydanov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Naydanov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Реализация поддержки диалектов в YaccConstructor/YARD',
         'author': 'Никонова',
         'supervisor': 'Кириленко', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Nikonova_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445-nikonova-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Восстановление адресного пространства процесса из расширенного образа памяти на платформе Windows',
         'author': 'Овчинников Антон Андреевич',
         'supervisor': 'Губанов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Ovchinnikov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Ovchinnikov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Ovchinnikov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка системы для мониторинга и анализа ботнетов, распространяемых через веб приложения',
         'author': 'Перевалова Марина Андреевна',
         'supervisor': 'Зеленчук', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Perevalova_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Perevalova_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Задача верификации лица на основе 3D модели',
         'author': 'Петров Николай Сергеевич',
         'supervisor': '', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Petrov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Petrov_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Petrov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Средства создания визуальных интерпретаторов диаграмм в системе QReal',
         'author': 'Поляков Владимир Александрович',
         'supervisor': 'Брыксин', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Polyakov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Polyakov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Polyakov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': ' Восстановление адресного пространства процесса из образа памяти с использованием файла подкачки на платформе Linux',
         'author': 'Свидерский Павел Юрьевич',
         'supervisor': 'Ãóáàíîâ', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Sviderski_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Sviderski_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Построение трёхмерной модели головы на основе трёхмерных анатомических признаков',
         'author': 'Серебряков Сергей Николаевич',
         'supervisor': 'Петров', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Serebryakov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Serebryakov_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Serebryakov_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Инструмент анализа пользовательских логов поисковых систем',
         'author': 'Солозобов Андрей Сергеевич',
         'supervisor': 'Грауэр', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Solozobov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Solozobov_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Solozobov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Система мониторинга веб-сервисов',
         'author': 'Фефелов Алексей Андреевич',
         'supervisor': 'Строкан', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Fefelov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Fefelov_review.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': '3D Registration',
         'author': 'Фоменко Екатерина Сергеевна',
         'supervisor': 'Пименов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Fomenko_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Fomenko_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Проектирование и реализация облачной metaCASE системы',
         'author': 'Чижова Надежда Александровна',
         'supervisor': 'Литвинов', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Chizhova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Chizhova_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Chizhova_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Практическая оценка качества различных средств HLS при синтезе из SystemC',
         'author': 'Шеин Роман Евгеньевич',
         'supervisor': 'Салищев', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Shein_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Shein_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Автоматическое тестирование верстки web-интерфейсов',
         'author': 'Шувалов Иннокентий Петрович',
         'supervisor': 'Ерошенко', 'publish_year': 2012,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Shuvalov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Shuvalov_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2012/YearlyProjects/2012/445/445_Shuvalov_review.pdf'
    )


def bruteforce_2011():
    # 345 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Синтаксический анализатор языка С',
         'author': 'Авдюхин Дмитрий Алексеевич',
         'supervisor': 'Кириленко', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Avdyukhin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Avdyukhin_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Avdyukhin_review.odt'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Организация надёжных соединений через виртуальные каналы',
         'author': 'Бурдун Фёдор Викторович',
         'supervisor': 'Бондарев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Burdun_report.pdf',
        slides_uri='',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Burdun_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Замена цвета выбранного элемента одежды в видеопотоке',
         'author': 'Григорьев Артем Валерьевич',
         'supervisor': 'Жуков', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Grigoryev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Grigoryev_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Grigoryev_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Графическая подсистема ОСРВ Embox для роботов LEGO Mindstorms NXT 2.0',
         'author': 'Дзендзик Дарья Анатольевна',
         'supervisor': 'Бондарев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Dzendzik_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Dzendzik_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Dzendzik_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Фрактальный анализ рынка',
         'author': 'Землянская Светлана Андреевна',
         'supervisor': 'Ширяев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Zemlyanskaya_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Zemlyanskaya_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Zemlyanskaya_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Алгоритмы вытеснения кэша. Применение алгоритмов в СХД и поиск возможностей их оптимизации для современных приложений.',
         'author': 'Колобов Роман Евгеньевич',
         'supervisor': 'Платонов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Kolobov_report.doc',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Kolobov_presentation.odp',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Помехоустойчивое свёрточное кодирование',
         'author': 'Коноплёв Юрий Михайлович',
         'supervisor': 'Татищев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Konoplyov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Konoplyov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Konoplyov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Оптимизация потоков для операционной системы Embox',
         'author': 'Мальковский Николай Владимирович',
         'supervisor': 'Бондарев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Malkovsky_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Malkovsky_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка декомпилятора языка Java SE 6',
         'author': 'Михайлов Дмитрий Петрович',
         'supervisor': 'Шафиров', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Mikhailov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Mikhailov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Mikhailov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Применение алгоритмов SuperResolution к лицам',
         'author': 'Мокаев Руслан',
         'supervisor': 'Пименов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Mokaev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Mokaev_presentation.pptx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Создание средств визуального сравнения моделей в QReal',
         'author': 'Мордвинов Дмитрий Александрович',
         'supervisor': 'Брыксин', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Mordvinov_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Mordvinov_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Аппаратное ускорение задачи выравнивания строк на языке HaSCoL',
         'author': 'Найданов Дмитрий Геннадьевич',
         'supervisor': 'Медведев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Naydanov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Naydanov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Naydanov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Предметно-ориентированное моделирование приложений для платформы Android',
         'author': 'Никонова Ольга Анатольевна',
         'supervisor': 'Брыксин', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Nikonova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Nikonova_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Nikonova_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Внедрение шифрования в систему хранения данных высокой производительности',
         'author': 'Овчинников Антон Андреевич',
         'supervisor': 'Ершов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Ovchinnikov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Ovchinnikov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Ovchinnikov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Определение положения камеры относительно плоского маркера',
         'author': 'Петров Николай',
         'supervisor': 'Вахитов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Petrov_report.doc',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Petrov_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/345/345-Petrov-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка визуального интерпретатора моделей в системе QReal',
         'author': 'Поляков Владимир Александрович',
         'supervisor': 'Брыксин', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Polyakov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Polyakov_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Polyakov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Обнаружение узлов в сети',
         'author': 'Свидерский Павел Юрьевич',
         'supervisor': 'Смирнов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Svidersky_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Svidersky_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Svidersky_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Восстановление положения объекта известной формы по зашумлённым наблюдениям с помощью видеокамеры',
         'author': 'Серебряков Сергей Николаевич',
         'supervisor': 'Вахитов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Serebryakov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Serebryakov_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345-Serebryakov-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Алгоритмы детектирующие и исправляющие ошибки в системах хранения данных уровня RAID 6',
         'author': 'Солозобов Андрей Сергеевич',
         'supervisor': 'Шевяков', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Solozobov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Solozobov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Solozobov_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Инструменты анализа данных метилирования генов в цепочке ДНК',
         'author': 'Фоменко Екатерина Сергеевна',
         'supervisor': 'Вяххи', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Fomenko_report.doc',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Fomenko_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Fomenko_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Библиотека алгоритмов поиска подстрок в тексте с препроцессингом в применении к биоинформатике',
         'author': 'Чижова Надежда Александровна',
         'supervisor': 'Вяххи', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Chizhova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Chizhova_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Chizhova_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация конфигурируемого аппаратного блока для вычисления быстрого преобразования Фурье переменной длины по смешанному основанию с использованием HLS',
         'author': 'Шеин Роман Евгеньевич',
         'supervisor': 'Салищев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Shein_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Shein_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Восстановление смазанных или размытых зашумлённых изображений',
         'author': 'Шувалов Иннокентий Петрович',
         'supervisor': 'Вахитов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Shuvalov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/345/345_Shuvalov_presentation.odp',
        supervisor_review_uri=''
    )
    # 361 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Аппаратное ускорение алгоритмов компьютерного зрения',
         'author': 'Стефан Бояровски',
         'supervisor': 'Шувалкин', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Boyarovski_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Boyarovski_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Boyarovski_review.rtf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Система проверки данных на полноту',
         'author': 'Карымов Антон Вячеславович',
         'supervisor': 'Графеев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Karymov_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Karymov_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Karymov_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Детектор элементов лица для построения псевдотрехмерной графики',
         'author': 'Лебедев Дмитрий',
         'supervisor': 'Пименов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Lebedev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Lebedev_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Публикация документов в SharePoint 2010',
         'author': 'Рябиченко Павел Николаевич',
         'supervisor': 'Дейкало', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Ryabichenko_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Ryabichenko_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Ryabichenko_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Разработка Eclipse плагина для программирования на языке haXe',
         'author': 'Савенко Мария Олеговна',
         'supervisor': 'Полозов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Savenko_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Savenko_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Savenko_review.odt'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Об оценке частоты устройств, разрабатываемых на языке HaSCoL',
         'author': 'Скородумов Кирилл Владимирович',
         'supervisor': 'Булычев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Skorodumov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Skorodumov_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Трансформация поисковых запросов в распределенных системах web-сервисов',
         'author': 'Солодка Анастасия Сергеевна',
         'supervisor': 'Новиков', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Solodkaya_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Solodkaya_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/361/361_Solodkaya_review.pdf'
    )
    # 445 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Улучшение автодополнения для языка Groovy в IDE IntelliJ IDEA',
         'author': 'Абишев Тимур Маратович',
         'supervisor': 'Мухин', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Abishev_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Abishev_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Abishev_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Программные инструменты и алгоритмы для определения информации о последовательности нуклеотидов по ее местоположению в геноме человека',
         'author': 'Алеев Алексей Валерьевич',
         'supervisor': 'Вяххи', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Aleev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Aleev_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Aleev_review.jpeg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Генерация объектной модели для DocsVision и использование ее при синхронизации сервисов',
         'author': 'Астащенко Александр Евгеньевич',
         'supervisor': 'Шистеров', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Astaschenko_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Astaschenko_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Astaschenko_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Семантическое автодополнение',
         'author': 'Василинец Сергей Павлович',
         'supervisor': 'Хитров', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Vasilinets_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Vasilinets_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Vasilinets_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Erlang. Статический и динамический анализ',
         'author': 'Гущина Вера Михайловна',
         'supervisor': 'Урбанович', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Gushchina_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Gushchina_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Gushchina_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Исправление краевых дефектов 3D скана лица и торса',
         'author': 'Добролеж Анна Борисовна',
         'supervisor': 'Антипов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Dobrolezh_report.doc',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Dobrolezh_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Dobrolezh_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Обработка структурных изменений источника данных в ETL процессах',
         'author': 'Долбешкин Андрей Николаевич',
         'supervisor': 'Дольник', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Dolbeshkin_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Dolbeshkin_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Dolbeshkin_review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Алгоритм приближённого join’а на потоках данных',
         'author': 'Землянский Юрий Андреевич',
         'supervisor': 'Новиков', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Zemlyanskiy_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Zemlyanskiy_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Zemlyanskiy_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка инструментария для создания искусственных нейронных сетей на мобильных платформах на примере iOS',
         'author': 'Золотухина Алина Игоревна',
         'supervisor': 'Торегожин', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Zolotukhina_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Zolotukhina_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Zolotukhina_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Обеспечение надёжности и высокой доступности кластера СХД',
         'author': 'Котов Юрий Александрович',
         'supervisor': 'Богатырев', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Kotov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Kotov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Kotov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Улучшение четкости и качества распознавания изображений роботов для системы SSL-Vision',
         'author': 'Кочанова Татьяна',
         'supervisor': 'Данилова', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Kochanova_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Kochanova_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Kochanova_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Сервер морфинга протеинов',
         'author': 'Лушников Андрей Сергеевич',
         'supervisor': 'Вяткина', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Lushnikov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Lushnikov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Lushnikov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Мотоциклетные графы и их свойства',
         'author': 'Мальчевский Михаил Андреевич',
         'supervisor': 'Вяткина', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Malchevsky_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Malchevsky_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка и апробация инструментария iOSNeuron для решения задач распознавания изображений',
         'author': 'Манаев Дмитрий Сергеевич',
         'supervisor': 'Торегожин', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Manayev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Manayev_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Manayev_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Исследование оптимизации запросов в СУБД',
         'author': 'Нишневич Анастасия Юрьевна',
         'supervisor': 'Новиков', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Nishnevich_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Nishnevich_presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Nishnevich_review.doc'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Метапрограммирование в .NET. Интерпретация Common Lisp',
         'author': 'Омельчук Александр Олегович',
         'supervisor': 'Полозов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Omelchuk_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Omelchuk_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Многоштриховые жесты мышью в проекте QReal',
         'author': 'Осечкина Мария Сергеевна',
         'supervisor': 'Литвинов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Osechkina_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Osechkina_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Osechkina_review.odt'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Нахождение сайтов начала репликации в ДНК человека',
         'author': 'Ромашкин Амир Сергеевич',
         'supervisor': 'Порозов', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Romashkin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Romashkin_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Romashkin_review.jpg'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Построение мотоциклетного графа',
         'author': 'Титов Артем Юрьевич',
         'supervisor': 'Вяткина', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Titov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Titov_presentation.ppt',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Titov_review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Семантическое автодополнение',
         'author': 'Удалов Александр Николаевич',
         'supervisor': 'Хитров', 'publish_year': 2011,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Udalov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Udalov_presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2011/YearlyProjects/2011/445/445_Udalov_review.pdf'
    )


def bruteforce_2010():
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Выявление семантических характеристик в слабоструктурированных текстовых данных',
         'author': 'Нурк Сергей Юрьевич',
         'supervisor': 'Вяткина', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Nurk_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Nurk_presentation.pdf',
        supervisor_review_uri=''
    )
    # 345 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Увеличение модульности программного обеспечения на языке Java',
         'author': 'Абишев Тимур Маратович',
         'supervisor': 'Сафонов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Abishev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Abishev_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Поддержка избыточного кодирования в проекте «Cirrostratus». Реализация алгоритмов избыточного кодирования на уровне ядра Linux',
         'author': 'Алеев Алексей Валерьевич',
         'supervisor': 'Косякин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Aleev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Aleev_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Автоматическая трансляция проекта Dypgen с языка OCaml на язык F#',
         'author': 'Баранов Эдуард Сергеевич',
         'supervisor': 'Кириленко', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Baranov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Baranov_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Translator Widget for Android',
         'author': 'Василинец Сергей Павлович',
         'supervisor': 'Филиппов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Vasilinets_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Vasilinets_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Анализ и постоение структуры сети',
         'author': 'Гущина Вера Михайловна',
         'supervisor': 'Никандров', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Gushina_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Gushina_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка масштабируемого интерфейса для клиентской программы с динамическим контентом',
         'author': 'Добролеж Анна Борисовна',
         'supervisor': 'Кириллин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Dobrolezh_report.doc',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Dobrolezh_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка редактора интерактивного электронного купона и реализация на платформе Flash',
         'author': 'Долбешкин Андрей Николаевич',
         'supervisor': 'Кириллин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Dolbeshkin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Dolbeshkin_presentation.pptx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация схемы распределённого поиска с использованием технологии Opera Unite',
         'author': 'Землянский Юрий Андреевич',
         'supervisor': 'Симуни', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Zemlyanskiy_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Zemlyanskiy_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Реализация PHP фреймворка',
         'author': 'Золотухина Алина Игоревна',
         'supervisor': 'Жуков', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Zolotukhina_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Zolotukhina_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Среда визуального моделирования on-line',
         'author': 'Иванов Всеволод Юрьевич',
         'supervisor': 'Литвинов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Ivanov_report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Ivanov_presentation.pptx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Система хранения данных. Поддержка избыточного кодирования. Поиск, сравнение и анализ применимости существующих подходов для поддержки избыточного кодирования',
         'author': 'Котов Юрий Александрович',
         'supervisor': 'Косякин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Kotov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Kotov_presentation.pptx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Поиск человека в красной куртке',
         'author': 'Кочанова',
         'supervisor': 'Вахитов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Kochanova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Kochanova_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка архитектуры многопоточного приложения под управлением операционной системы iPhone OS 2.2.1 и выше',
         'author': 'Лушников Андрей Сергеевич',
         'supervisor': 'Кириллин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Lushnikov_report.doc',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Lushnikov_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Система хранения данных. Поддержка избыточного кодирования. Оптимизация, настройка и апробация выбранного алгоритма под поставленную задачу. Оценка полученных результатов',
         'author': 'Мальчевский Михаил Андреевич',
         'supervisor': 'Косякин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Malchevsky_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Malchevsky_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Внедрение unit-тестирования в проект на F#',
         'author': 'Нишневич Анастасия Юрьевна',
         'supervisor': 'Кириленко', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Nishnevich_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Nishnevich_presentation.pptx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Базовые алгоритмы файлового карвинга',
         'author': 'Омельчук Александр Олегович',
         'supervisor': 'Губанов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Omelchuk_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Omelchuk_presentation.ppsx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Визуальное программирование при помощи мыши',
         'author': 'Осечкина Мария Сергеевна',
         'supervisor': 'Литвинов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Osechkina_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Osechkina_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5, 'name_ru': 'Распознавание движения человека по ряду изображений',
         'author': 'Ромашкин Амир Сергеевич',
         'supervisor': 'Вахитов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Romashkin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Romashkin_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Обучающая программа для медицинского тренажера «Максим»',
         'author': 'Титов',
         'supervisor': 'Дубчук', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Titov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Titov_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка архитектуры сетевого многопоточного приложения системы мобильного маркетинга на платформе Java ME',
         'author': 'Удалов Александр Николаевич',
         'supervisor': 'Кириллин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Udalov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Udalov_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Сетевой анализатор трафика на языке программирования Perl в ОС Windows',
         'author': 'Филиппова Анастасия Валерьевна',
         'supervisor': 'Баклановский', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Filippova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Filippova_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка среды для облачных вычислений',
         'author': 'Чуновкин Фёдор Дмитриевич',
         'supervisor': 'Бондарев', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Chunovkin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/345/Chunovkin_presentation.pptx',
        supervisor_review_uri=''
    )
    # 361 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Перенос драйвера блочного устройства DST на уровень Ethernet для проекта Cirrostratus',
         'author': 'Колянов Дмитрий Андреевич',
         'supervisor': 'Богатырев', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Kolyanov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Kolyanov_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Анализ, сравнение и адаптация протоколов для оптимальной передачи данных в проекте Cirrostratus',
         'author': 'Кузнецов Кирилл Олегович',
         'supervisor': 'Богатырев', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Kuznetsov_report.odt',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Kuznetsov_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Разработка надежного протокола обмена данными на уровне Ethernet в проекте Cirrostratus',
         'author': 'Лапин Сергей Константинович',
         'supervisor': 'Богатырев', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Lapin_report.odt',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Lapin_presentation.odp',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Реализация алгоритма минимизации стоимости потребления электроэнергии',
         'author': 'Цыпан Ксения Владимировна',
         'supervisor': 'Графеева', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Tsipan_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/361/Tsipan_presentation.ppt',
        supervisor_review_uri=''
    )
    # 445 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка системы тестирования программно-аппаратных комплексов',
         'author': 'Батюков Александр Михайлович',
         'supervisor': 'Бондарев', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Batyukov_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка метода сбора информации о ходе исполнения программы, который использует возможность модификации памяти процесса',
         'author': 'Булычев Иван Дмитриевич',
         'supervisor': 'Баклановский', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Boulichev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Boulichev_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Оптимизация процесса сборки документов системой nutch',
         'author': 'Волков Сергей Андреевич',
         'supervisor': 'Выговский', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Volkov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Volkov_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация мобильных сервисов для доступа к удаленным устройствам на базе платформы Ubiq Mobile',
         'author': 'Гладышева Юлия Сергеевна',
         'supervisor': 'Оносовский', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Gladisheva_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Gladisheva_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка средства для передачи информации через экран мобильного устройства',
         'author': 'Дьяченко Василий Владимирович',
         'supervisor': 'Кириллин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Dyachenko_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Dyachenko_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка сервера системы мобильного маркетинга',
         'author': 'Зарубин Михаил Сергеевич',
         'supervisor': 'Кириллин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Zarubin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Zarubin_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Совместимость филогенетических деревьев',
         'author': 'Катышев Алексей Александрович',
         'supervisor': 'Вяткина', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Katyshev_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Создание системы хранения и выдачи данных',
         'author': 'Константинов Александр Сергеевич',
         'supervisor': 'Лопатин', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Konstantinov_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Konstantinov_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Псевдо-треугольники и псевдо-четырехугольники с пустой внутренностью',
         'author': 'Копелиович Сергей Владимирович',
         'supervisor': 'Вяткина', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Kopeliovich_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Kopeliovich_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Система отслеживания документов на письменном столе',
         'author': 'Кривоконь Дмитрий Сергеевич',
         'supervisor': 'Вахитов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Krivokon_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Krivokon_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Поиск шаблонов в программном коде',
         'author': 'Куделевский Евгений Валерьевич',
         'supervisor': 'Мосиенко', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Kudelevsky_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Kudelevsky_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Кроссъязыковый рефакторинг «Изменение сигнатуры метода» для IDE IntelliJ IDEA',
         'author': 'Медведев Максим Юрьевич',
         'supervisor': 'Громов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Medvedev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Medvedev_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Сегментация речи по источнику первичного возбуждения, определение артикуляционных классов сегментов',
         'author': 'Меламуд Александр Евгеньевич',
         'supervisor': 'Булашевич', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Melamud_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Melamud_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Реализация подхода Scrap Your Boilerplate для Ocaml',
         'author': 'Мечтаев Сергей Владимирович',
         'supervisor': 'Булычев', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Mechtaev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Mechtaev_presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация субпиксельного уточнения ViFlow метода поиска оптического потока',
         'author': 'Расторгуев Алексей Сергеевич',
         'supervisor': 'Пименов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Rastorguyev_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Rastorguyev_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Интеграция CASE-системы QReal с Scilab',
         'author': 'Савин Никита Сергеевич',
         'supervisor': 'Литвинов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Savin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Savin_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Анализ применимости семантического кеширования на основе подобия',
         'author': 'Анна Сафонова',
         'supervisor': 'Новиков', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Safonova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Safonova_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Определение импульса основного тона сигнала с плохим соотношением «сигнал-шум»',
         'author': 'Такун Евгения Игоревна',
         'supervisor': 'Булашевич', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Takun_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Takun_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Модель и алгоритм улучшения распознавания частей речи в текстах, содержащих ошибки',
         'author': 'Ткаченко Максим Владиславович',
         'supervisor': 'Выговский', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Tkachenko_report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация мобильных сервисов для доступа к удаленным устройствам на базе платформы Ubiq Mobile',
         'author': 'Туманова Кристина Сергеевна',
         'supervisor': 'Оносовский', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Tumanova_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Tumanova_presentation.ppt',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка архитектуры для генератора синтаксических анализаторов',
         'author': 'Улитин Константин Андреевич',
         'supervisor': 'Кириленко', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Ulitin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Ulitin_presentation.pptx',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Особенности open-source разработки на основе проекта “embox”',
         'author': 'Фомин Алексей Дмитриевич',
         'supervisor': 'Бондарев', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Fomin_report.doc',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Карвинг сжатых NTFS разделов',
         'author': 'Щитинин Дмитрий Анатольевич',
         'supervisor': 'Губанов', 'publish_year': 2010,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Shitinin_report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2010/YearlyProjects/2010/445/Shitinin_presentation.ppt',
        supervisor_review_uri=''
    )


def bruteforce_2009():
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Параллельная реализация алгоритма ACO',
         'author': 'Дырдина Анна Викторовна',
         'supervisor': 'Вахитов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Dyrdina_Anna.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Интеграция технологии DocLine с системой разработки документации Adobe',
         'author': 'Дорохов Вадим Александрович',
         'supervisor': 'Романовский', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Dorokhov_Vadim.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Интеграция мультимедиа решений с аппаратным ускорением для MID устройства',
         'author': 'Елизаров Егор Алексеевич',
         'supervisor': 'Иванов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Elisarov_Egor.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Конвертор приложений Windows Forms в приложения Silverlight',
         'author': 'Гагарский Алексей Константинович',
         'supervisor': 'Губанов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Gagarsky_Alexey.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Расширение функциональности Entity Framework',
         'author': 'Хритошин Даниил Викторович',
         'supervisor': 'Евдокимов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Khritoshin_Daniil.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Создание среды разработки библиотек формул подсчета технико-экономических показателей теплоэлектростанций',
         'author': 'Иноземцев Дмитрий Сергеевич',
         'supervisor': 'Иванов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Inozemtsev_Dmitry.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Массовая задача построения маршрутов движения судов',
         'author': 'Кудасов Федор Сергеевич',
         'supervisor': 'Кариженский', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Kudasov_Fyodor.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Автоматизация отслеживания состояния покрытия автомобильных дорог. Использование мобильных устройств, оснащённых акселерометром и устройством определения местоположения для определения состояния дорожного покрытия и для обнаружения дефектов дорожного полотна',
         'author': 'Леви Сергей',
         'supervisor': 'Сабашный', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Levi_Sergei.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка SIP телефонии для операционной системы Google Android',
         'author': 'Малышев Виталий',
         'supervisor': 'Сафонов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Malyshev_Vitaly.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка приложения для платформы Google Android',
         'author': 'Морозков Михаил Андреевич',
         'supervisor': 'Замышляев', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Morozkov_Mikhail.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка расширения языка Java для работы с реляционными базами данных',
         'author': 'Никитин Павел Антонович',
         'supervisor': 'Гуров', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Nikitin_Pavel.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Декомпозиция временных рядов в СУБД Oracle',
         'author': 'Подкорытов Сергей',
         'supervisor': 'Графеева', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Podkorytov_Sergei.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Создание языка для проверки свойств контекстно-свободных грамматик',
         'author': 'Силина Ольга Александровна',
         'supervisor': 'Бреслав', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Silina_Olga.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Библиотека алгоритмов на графах для платформы .NET',
         'author': 'Суханов Василий',
         'supervisor': 'Кириленко', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Sukhanov_Vasily.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Разработка программного обеспечения системы программно-аппаратной защиты ПО',
         'author': 'Теплых Дарья Анатольевна',
         'supervisor': 'Татищев', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Teplyh_Daria.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Расширение функциональности графического редактора языка DRL',
         'author': 'Василик Дмитрий',
         'supervisor': 'Романовский', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Vasilik_Dmitry.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'C++ APPLICATION SECURITY TOOLSET',
         'author': 'Ростислав Игоревич Чутков',
         'supervisor': 'Штукенберг', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/444_Chutkov_Rostislav.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    # 445 группа
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Поддержка языка Lisa в среде Eclipse',
         'author': 'Алеев Константин',
         'supervisor': 'Иванов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Aleev_Konstantin.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Инструмент аспектно-ориентированного программирования Aspect.Java',
         'author': 'Андриевский Евгений Валерьевич',
         'supervisor': 'Сафонов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Andrievsky_Evgeny.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Разработка framework для JSR 290 TCK',
         'author': 'Евстифеев Сергей Викторович',
         'supervisor': 'Исаенко', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Evstifeev_Sergei.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Распознавание автомобильных номеров с помощью нейронных сетей',
         'author': 'Федяшов Виктор',
         'supervisor': 'Пименов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Fedyashov_Victor.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Создание дискретизирующего фильтра для обработки электроокулограмм. Повышение точности определения амплитуды сигнала',
         'author': 'Храмцова Елена Александровна',
         'supervisor': 'Белов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Khramtsova_Elena.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Реализация подключения виртуальной машины Neko к http-серверу с помощью интерфейса FastCGI',
         'author': 'Ларчик Евгений Владимирович',
         'supervisor': 'Плискин', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Larchik_Evgeny.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Создание среды разработки для языка программирования OCaml',
         'author': 'Мануйлов Максим Игоревич',
         'supervisor': 'Шафиров', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Manuilov_Maxim.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Cоздание дискретизирующего фильтра для обработки электроокулограмм. Обеспечение работы и настройки фильтра в режиме реального времени',
         'author': 'Медведев Алексей Михайлович',
         'supervisor': 'Белов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Medvedev_Alexey.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Язык для описания плагинов в среде программирования JetBrains MPS',
         'author': 'Мухин Михаил',
         'supervisor': 'Соломатов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Mukhin_Mikhail.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Применение нейронных сетей к ранжированию результатов информационного поиска',
         'author': 'Петров Александр Георгиевич',
         'supervisor': 'Вахитов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Petrov_Alexander.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Поиск оптимального ректификационного преобразования',
         'author': 'Смирнова Ольга',
         'supervisor': 'Пименов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Smirnova_Olga.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Акторное расширение языка Java в среде MPS',
         'author': 'Жукова Анна Руслановна',
         'supervisor': 'Мазин', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Smirnova_Olga.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Исследование работы с географическими данными в Oracle 10g',
         'author': 'Залог Леонид Витальевич',
         'supervisor': 'Графеева', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Zalog_Leonid.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Создание физически-корректного дождя и сопутствующих эффектов',
         'author': 'Шевченко Андрей Игоревич',
         'supervisor': 'Пименов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Shevchenko_Andrei.7z',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1, 'name_ru': 'Модульная платформа для распознования автомобильных номеров',
         'author': 'Чирков Иван Викторович',
         'supervisor': 'Пименов', 'publish_year': 2009,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2009/YearlyProjects/2009/445_Chirkov_Ivan.7z',
        slides_uri='',
        supervisor_review_uri=''
    )


def bruteforce_2015_fall():
    upload_one_report(
        {'type_id': 2, 'course_id': 3,
         'name_ru': 'Построение дискретного плана управления по спецификации в виде LTL-формулы',
         'author': 'Агапова Татьяна Юрьевна',
         'supervisor': 'Литвинов', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/autumn-2015/magistracy-564/agapova-tatyana-yurevna/at_download/file',
        slides_uri='',
        supervisor_review_uri=''
    )


def bruteforce_2015_spring():
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Дешифрация образа диска с защитой Bitlocker To Go инструментами анализа дампа памяти',
         'author': 'Грабовой Филипп Николаевич',
         'supervisor': 'Губанов', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Grabovoy-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Grabovoy-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Grabovoy-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Система автоматизированного массового тестирования проекта CODA',
         'author': 'Комаров Константин Михайлович',
         'supervisor': 'Баклановский', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Komarov-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Komarov-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Komarov-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Автоматическое тестирование пользовательского интерфейса системы QReal',
         'author': 'Никольский Кирилл Андреевич',
         'supervisor': 'Литвинов', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Nikolskiy-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Nikolskiy-presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Создание системы проектирования БД на базе платформы QReal',
         'author': 'Семенова Анастасия Владимировна',
         'supervisor': 'Брыксин', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Semenova-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Semenova-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Semenova-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Система мониторинга MSSQL сервера: отслеживание и анализ характеристик производительности',
         'author': 'Столпнер Лев Артемович',
         'supervisor': 'Давыденко', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Stolpner-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Stolpner-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Stolpner-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Распознавание характеристик объектов в робофутболе',
         'author': 'Черняев Арсений Витальевич',
         'supervisor': 'Пименов', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Chernyaev-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/344/344-Chernyaev-presentation.pdf',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Анимирование речи',
         'author': 'Брыксин Матвей Александрович',
         'supervisor': 'Вахитов', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/371/371-Bryksin-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/371/371-Bryksin-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/371/371-Bryksin-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 2,
         'name_ru': 'Интерпретация диаграмм в онлайн-среде программирования роботов',
         'author': 'Гагина Лада Владиславовна',
         'supervisor': 'Брыксин', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/371/371-Gagina-report.pdf',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 2,
         'name_ru': 'Метрика в пространстве портретов процессов',
         'author': 'Лозов Петр Алексеевич',
         'supervisor': 'Баклановский', 'publish_year': 2015,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/371/371-Lozov-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/371/371-Lozov-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/spring-2015/371/371-Lozov-review.pdf'
    )


def bruteforce_2013():
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Методы взаимодействия прикладного приложения и ядра ОС',
         'author': 'Булычев Антон Дмитриевич',
         'supervisor': 'Абусалимов', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/344/344-Bulychev-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/344/344-Bulychev-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/344/344-Bulychev-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Декомпиляция выражений по байт-коду JVM',
         'author': 'Поздин Дмитрий Евгеньевич',
         'supervisor': 'Булычев', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/344/344-Pozdin-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/344/344-Pozdin-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/344/344-Pozdin-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Метод реконструкции невидимых областей полигональных 3D моделей',
         'author': 'Егорова Елизавета Сергеевна',
         'supervisor': 'Петров', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/361/361-Egorova-report.docx',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/361/361-Egorova-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/361/361-Egorova-review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 5,
         'name_ru': 'Параметризация и изменение формы 3D модели человека',
         'author': 'Тарасова Евгения Сергеевна',
         'supervisor': 'Петров', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/361/361-Tarasova-report.doc',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/361/361-Tarasova-presentation.pptx',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/361/361-Tarasova-review.docx'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Cетевой стек реального времени',
         'author': 'Калмук Александр Игоревич',
         'supervisor': 'Абусалимов', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Kalmuk-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Kalmuk-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Kalmuk-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Форматирование текста программ на основе комбинаторов, сопоставления с образцом и синтаксических шаблонов',
         'author': 'Подкопаев Антон Викторович',
         'supervisor': 'Булычев', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Podkopaev-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Podkopaev-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Podkopaev-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': 'Использование proof assistants для описания операционных семантик',
         'author': 'Таран Кирилл Сергеевич',
         'supervisor': 'Булычев', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Taran-report.pdf',
        slides_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Taran-presentation.pdf',
        supervisor_review_uri='https://oops.math.spbu.ru/SE/YearlyProjects/2013/YearlyProjects/2013/445/445-Taran-review.pdf'
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': '',
         'author': '',
         'supervisor': '', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': '',
         'author': '',
         'supervisor': '', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='',
        slides_uri='',
        supervisor_review_uri=''
    )
    upload_one_report(
        {'type_id': 2, 'course_id': 1,
         'name_ru': '',
         'author': '',
         'supervisor': '', 'publish_year': 2013,
         'secret_key': SECRET_KEY},
        text_uri='',
        slides_uri='',
        supervisor_review_uri=''
    )


if __name__ == '__main__':
    bruteforce_2013()
    bruteforce_2015_spring()
    bruteforce_2015_fall()
    # bruteforce_2009()
    # bruteforce_2010()
    # bruteforce_2011()
    # bruteforce_2012()
    # get_2013()
    # get_2014()
    # get_2015_spring()
    # get_2015_fall()
    # get_2016_reports()
    # get_2017_reports()
